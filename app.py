
import os, sqlite3, re, smtplib, ssl
from flask import Flask, render_template, request, redirect, session, url_for, flash, send_from_directory, abort, jsonify, make_response
from werkzeug.security import generate_password_hash, check_password_hash
from datetime import datetime
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from io import BytesIO
from docx import Document
from openpyxl import Workbook
import stripe

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DB_PATH = os.path.join(BASE_DIR, 'database.db')
UPLOAD_FOLDER = os.path.join(BASE_DIR, 'uploads')
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

app = Flask(__name__)
app.secret_key = os.environ.get('SECRET_KEY', 'change_this_to_secret')
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

EMAIL_RE = re.compile(r"[^@]+@[^@]+\.[^@]+")
# TEST MODE KEYS (set in Render env)
stripe.api_key = os.environ.get('STRIPE_SECRET_KEY', 'sk_test_your_key_here')
STRIPE_PUBLIC_KEY = os.environ.get('STRIPE_PUBLIC_KEY', 'pk_test_your_key_here')
STRIPE_CURRENCY = os.environ.get('STRIPE_CURRENCY', 'rub')
STRIPE_WEBHOOK_SECRET = os.environ.get('STRIPE_WEBHOOK_SECRET', 'whsec_your_webhook_secret')

def get_db():
    conn = sqlite3.connect(DB_PATH); conn.row_factory = sqlite3.Row; return conn

def init_db():
    with open(os.path.join(BASE_DIR,'models.sql'),'r',encoding='utf-8') as f: sql=f.read()
    conn=get_db(); conn.executescript(sql)
    cur=conn.cursor()
    if cur.execute("SELECT COUNT(*) FROM roles").fetchone()[0]==0:
        cur.execute("INSERT INTO roles (name) VALUES ('admin'),('manager'),('client')")
    # seed users
    def ensure_user(u,e,pw,role):
        if cur.execute("SELECT COUNT(*) FROM users WHERE username=?",(u,)).fetchone()[0]==0:
            cur.execute("INSERT INTO users (username,email,password,role_id) VALUES (?,?,?,(SELECT id FROM roles WHERE name=?))",(u,e,generate_password_hash(pw),role))
    ensure_user('admin','admin@example.com','AdminPass123','admin')
    ensure_user('manager','manager@example.com','ManagerPass123','manager')
    ensure_user('student','student@witte.edu.ru','Student123','client')
    if cur.execute("SELECT COUNT(*) FROM tariffs").fetchone()[0]==0:
        cur.executemany("INSERT INTO tariffs (name,description,price) VALUES (?,?,?)", [
            ('Очная (бакалавриат)','Оплата семестра очной формы обучения',120000),
            ('Заочная (бакалавриат)','Оплата семестра заочной формы обучения',70000),
            ('Магистратура (очная)','Оплата семестра магистратуры',150000),
            ('Колледж (очная)','Оплата семестра колледжа',65000),
        ])
    conn.commit(); conn.close(); print('✅ Database initialized:', DB_PATH)

if not os.path.exists(DB_PATH): init_db()

def current_user():
    if 'user_id' in session:
        conn=get_db(); u=conn.execute("SELECT users.*, roles.name role_name FROM users LEFT JOIN roles ON users.role_id=roles.id WHERE users.id=?",(session['user_id'],)).fetchone(); conn.close(); return u
    return None

def login_required(f):
    from functools import wraps
    @wraps(f)
    def inner(*a,**kw):
        if 'user_id' not in session: return redirect(url_for('login', next=request.path))
        return f(*a,**kw)
    return inner

def role_required(*roles):
    from functools import wraps
    def dec(f):
        @wraps(f)
        def inner(*a,**kw):
            u=current_user()
            if not u or u['role_name'] not in roles:
                flash('Доступ запрещён'); return redirect(url_for('index'))
            return f(*a,**kw)
        return inner
    return dec

def send_email_with_pdf(to_email, subject, html_body, pdf_path, cc=None):
    host=os.environ.get('SMTP_HOST','smtp.mailtrap.io'); port=int(os.environ.get('SMTP_PORT',587))
    user=os.environ.get('SMTP_USER','user'); password=os.environ.get('SMTP_PASS','pass')
    sender=os.environ.get('SMTP_FROM','no-reply@vitte.example')
    msg=MIMEMultipart(); msg['From']=sender; msg['To']=to_email
    if cc: msg['Cc']=cc
    msg['Subject']=subject; msg.attach(MIMEText(html_body,'html','utf-8'))
    if pdf_path and os.path.exists(pdf_path):
        with open(pdf_path,'rb') as f: part=MIMEApplication(f.read(),_subtype='pdf')
        part.add_header('Content-Disposition','attachment', filename=os.path.basename(pdf_path)); msg.attach(part)
    ctx=ssl.create_default_context()
    with smtplib.SMTP(host,port) as s:
        try:
            s.starttls(context=ctx); s.login(user,password)
        except Exception as e:
            print('SMTP note:', e)
        s.send_message(msg)

def make_receipt_pdf(payment_id, fio, program, amount, currency):
    pdf_name=f"receipt_{payment_id}.pdf"; pdf_path=os.path.join(UPLOAD_FOLDER,pdf_name)
    c=canvas.Canvas(pdf_path, pagesize=A4); w,h=A4
    c.setFont('Helvetica-Bold',14); c.drawString(50,h-50,'ЧОУВО «Московский университет имени С.Ю. Витте»')
    c.setFont('Helvetica',11); c.drawString(50,h-80,'Квитанция об оплате обучения'); c.line(50,h-85,w-50,h-85)
    y=h-120
    for line in [f'Номер платежа: {payment_id}', f'ФИО: {fio or "-"}', f'Программа: {program}', f'Сумма: {amount:.2f} {currency.upper()}', f'Дата: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}']:
        c.drawString(50,y,line); y-=20
    c.showPage(); c.save()
    return pdf_name, pdf_path

@app.context_processor
def inject_user():
    return dict(current_user=current_user)

@app.route('/')
def index():
    conn=get_db(); tariffs=conn.execute("SELECT * FROM tariffs ORDER BY price DESC").fetchall(); conn.close()
    return render_template('index.html', tariffs=tariffs)

@app.route('/tuition')
def tuition():
    conn=get_db(); tariffs=conn.execute("SELECT * FROM tariffs ORDER BY price DESC").fetchall(); conn.close()
    return render_template('tuition.html', tariffs=tariffs)

@app.route('/faq') 
def faq(): return render_template('faq.html')

@app.route('/contact', methods=['GET','POST'])
def contact():
    if request.method=='POST':
        name=request.form.get('name','').strip(); email=request.form.get('email','').strip(); msg=request.form.get('message','').strip()
        if not name or not email or not msg or not EMAIL_RE.match(email): flash('Проверьте поля'); return redirect(url_for('contact'))
        conn=get_db(); conn.execute("INSERT INTO contacts (name,email,message) VALUES (?,?,?)",(name,email,msg)); conn.commit(); conn.close()
        flash('Сообщение отправлено'); return redirect(url_for('contact'))
    return render_template('contact.html')

@app.route('/register', methods=['GET','POST'])
def register():
    if request.method=='POST':
        username=request.form.get('username','').strip(); email=request.form.get('email','').strip(); password=request.form.get('password','').strip()
        if not username or not email or not password or not EMAIL_RE.match(email): flash('Проверьте поля'); return redirect(url_for('register'))
        conn=get_db()
        try:
            conn.execute("INSERT INTO users (username,email,password,role_id) VALUES (?,?,?,(SELECT id FROM roles WHERE name='client'))",(username,email,generate_password_hash(password))); conn.commit(); flash('Регистрация успешна. Войдите.'); return redirect(url_for('login'))
        except sqlite3.IntegrityError: flash('Логин занят')
        finally: conn.close()
    return render_template('register.html')

@app.route('/login', methods=['GET','POST'])
def login():
    if request.method=='POST':
        username=request.form.get('username','').strip(); password=request.form.get('password','').strip()
        conn=get_db(); u=conn.execute("SELECT * FROM users WHERE username=?",(username,)).fetchone(); conn.close()
        if u and check_password_hash(u['password'], password): session['user_id']=u['id']; session['username']=u['username']; flash('Вход выполнен'); return redirect(url_for('index'))
        else: flash('Неверный логин или пароль')
    return render_template('login.html')

@app.route('/logout')
def logout(): session.clear(); flash('Вы вышли'); return redirect(url_for('login'))

@app.route('/checkout/<int:tariff_id>')
def checkout(tariff_id):
    conn=get_db(); t=conn.execute("SELECT * FROM tariffs WHERE id=?",(tariff_id,)).fetchone(); conn.close()
    if not t: abort(404)
    currency=os.environ.get('STRIPE_CURRENCY','rub'); unit_amount=int(round(float(t['price'])*100))
    session_obj=stripe.checkout.Session.create(
        mode='payment', payment_method_types=['card'],
        line_items=[{'price_data':{'currency':currency,'unit_amount':unit_amount,'product_data':{'name':t['name']}},'quantity':1}],
        success_url=url_for('success',_external=True)+"?session_id={CHECKOUT_SESSION_ID}",
        cancel_url=url_for('cancel',_external=True),
        metadata={'tariff_id':str(tariff_id),'tariff_name':t['name']}
    )
    conn=get_db()
    conn.execute("INSERT INTO payments (user_id, fio, program, amount, currency, status, stripe_session_id) VALUES (?,?,?,?,?,?,?)",
                 (session.get('user_id'), None, t['name'], float(t['price']), currency, 'pending', session_obj.id))
    conn.commit(); conn.close()
    return render_template('checkout.html', public_key=STRIPE_PUBLIC_KEY, checkout_session_id=session_obj.id, tariff=t)

@app.route('/success')
def success():
    sess_id=request.args.get('session_id'); receipt_link=None; name=''
    if sess_id:
        conn=get_db(); p=conn.execute("SELECT * FROM payments WHERE stripe_session_id=?",(sess_id,)).fetchone()
        if p: name=p['program']
        if p and p['receipt_pdf']: receipt_link=url_for('download_receipt', payment_id=p['id'])
        conn.close()
    return render_template('success.html', tariff_name=name, receipt_link=receipt_link)

@app.route('/cancel')
def cancel(): return render_template('cancel.html')

@app.route('/stripe/webhook', methods=['POST'])
def stripe_webhook():
    payload=request.data; sig=request.headers.get('Stripe-Signature')
    try:
        event=stripe.Webhook.construct_event(payload, sig, STRIPE_WEBHOOK_SECRET)
    except Exception as e:
        return jsonify({'status':'invalid','error':str(e)}), 400

    if event['type']=='checkout.session.completed':
        s=event['data']['object']; sess_id=s['id']
        conn=get_db(); p=conn.execute("SELECT * FROM payments WHERE stripe_session_id=?",(sess_id,)).fetchone()
        if p and p['status']!='paid':
            conn.execute("UPDATE payments SET status='paid' WHERE id=?",(p['id'],))
            pdf_name, pdf_path = make_receipt_pdf(p['id'], p['fio'], p['program'], p['amount'], p['currency'])
            conn.execute("UPDATE payments SET receipt_pdf=? WHERE id=?", (pdf_name, p['id'])); conn.commit()
            # email
            email=None; username='Студент'
            if p['user_id']:
                u=conn.execute("SELECT * FROM users WHERE id=?",(p['user_id'],)).fetchone()
                if u: email=u['email']; username=u['username']
            conn.close()
            html=render_template('emails/receipt.html', fio=p['fio'], username=username, program=p['program'], amount=p['amount'], currency=p['currency'], created_at=p['created_at'], receipt_no=p['id'])
            billing=os.environ.get('SMTP_TO_BILLING')
            try:
                if email: send_email_with_pdf(email,'Квитанция об оплате обучения',html,pdf_path,cc=billing)
                elif billing: send_email_with_pdf(billing,'Новый платёж за обучение',html,pdf_path)
            except Exception as e:
                print('Email error:',e)
        else:
            conn.close()
    return jsonify({'status':'ok'})

@app.route('/receipts/<int:payment_id>.pdf')
@login_required
def download_receipt(payment_id):
    conn=get_db(); p=conn.execute("SELECT * FROM payments WHERE id=? AND user_id=?",(payment_id,session.get('user_id'))).fetchone(); conn.close()
    if not p or not p['receipt_pdf']: abort(404)
    return send_from_directory(app.config['UPLOAD_FOLDER'], p['receipt_pdf'], as_attachment=True)

@app.route('/dashboard')
@login_required
def dashboard():
    conn=get_db(); pays=conn.execute("SELECT * FROM payments WHERE user_id=? ORDER BY created_at DESC",(session['user_id'],)).fetchall(); conn.close()
    return render_template('user/dashboard.html', payments=pays)

@app.route('/dashboard/profile', methods=['GET','POST'])
@login_required
def profile():
    u=current_user()
    if request.method=='POST':
        email=request.form.get('email','').strip(); password=request.form.get('password','').strip()
        conn=get_db()
        if password: conn.execute("UPDATE users SET email=?, password=? WHERE id=?", (email, generate_password_hash(password), u['id']))
        else: conn.execute("UPDATE users SET email=? WHERE id=?", (email, u['id']))
        conn.commit(); conn.close(); flash('Профиль обновлён'); return redirect(url_for('profile'))
    return render_template('user/profile.html', user=u)

@app.route('/dashboard/documents')
@login_required
def user_documents():
    conn=get_db(); docs=conn.execute("SELECT id, program, amount, created_at, receipt_pdf FROM payments WHERE user_id=? ORDER BY created_at DESC",(session['user_id'],)).fetchall(); conn.close()
    return render_template('user/documents.html', docs=docs)

# ===== Admin =====
def stats_data():
    conn=get_db()
    total=conn.execute("SELECT COUNT(*) c FROM payments").fetchone()['c']
    paid=conn.execute("SELECT COUNT(*) c FROM payments WHERE status='paid'").fetchone()['c']
    summ=conn.execute("SELECT SUM(amount) s FROM payments WHERE status='paid'").fetchone()['s'] or 0
    conn.close()
    return {'total_count':total,'paid_count':paid,'total_sum':summ}

@app.route('/admin')
@role_required('admin','manager')
def admin_index(): return render_template('admin/admin_index.html')

@app.route('/admin/stats')
@role_required('admin','manager')
def admin_stats(): return render_template('admin/stats.html', stats=stats_data())

@app.route('/admin/tariffs')
@role_required('admin')
def admin_tariffs():
    conn=get_db(); tariffs=conn.execute("SELECT * FROM tariffs ORDER BY price DESC").fetchall(); conn.close()
    return render_template('admin/tariffs.html', tariffs=tariffs)

@app.route('/admin/tariffs/create', methods=['GET','POST'])
@role_required('admin')
def admin_tariff_create():
    if request.method=='POST':
        name=request.form.get('name','').strip(); description=request.form.get('description','').strip(); price=float(request.form.get('price','0').strip())
        conn=get_db(); conn.execute("INSERT INTO tariffs (name,description,price) VALUES (?,?,?)",(name,description,price)); conn.commit(); conn.close(); flash('Тариф создан'); return redirect(url_for('admin_tariffs'))
    return render_template('admin/tariff_form.html', title='Новый тариф', t=None)

@app.route('/admin/tariffs/<int:tid>/edit', methods=['GET','POST'])
@role_required('admin')
def admin_tariff_edit(tid):
    conn=get_db(); t=conn.execute("SELECT * FROM tariffs WHERE id=?",(tid,)).fetchone()
    if not t: conn.close(); abort(404)
    if request.method=='POST':
        name=request.form.get('name','').strip(); description=request.form.get('description','').strip(); price=float(request.form.get('price','0').strip())
        conn.execute("UPDATE tariffs SET name=?, description=?, price=? WHERE id=?", (name,description,price,tid)); conn.commit(); conn.close(); flash('Тариф обновлён'); return redirect(url_for('admin_tariffs'))
    conn.close(); return render_template('admin/tariff_form.html', title='Редактирование тарифа', t=t)

@app.route('/admin/tariffs/<int:tid>/delete', methods=['POST'])
@role_required('admin')
def admin_tariff_delete(tid):
    conn=get_db(); conn.execute("DELETE FROM tariffs WHERE id=?", (tid,)); conn.commit(); conn.close(); flash('Удалено'); return redirect(url_for('admin_tariffs'))

@app.route('/admin/payments')
@role_required('admin','manager')
def admin_payments():
    conn=get_db(); pays=conn.execute("SELECT * FROM payments ORDER BY created_at DESC").fetchall(); conn.close()
    return render_template('admin/payments.html', payments=pays)

@app.route('/admin/users')
@role_required('admin')
def admin_users():
    conn=get_db(); users=conn.execute("SELECT users.*, roles.name role_name FROM users LEFT JOIN roles ON users.role_id=roles.id ORDER BY users.id DESC").fetchall(); conn.close()
    return render_template('admin/users.html', users=users)

# Exports
@app.route('/admin/exports')
@role_required('admin')
def admin_exports():
    return render_template('admin/exports.html')

@app.route('/admin/export/xlsx')
@role_required('admin')
def export_xlsx():
    conn=get_db(); rows=conn.execute("SELECT id, user_id, fio, program, amount, currency, status, created_at FROM payments ORDER BY created_at DESC").fetchall(); conn.close()
    wb=Workbook(); ws=wb.active; ws.append(['ID','User ID','FIO','Program','Amount','Currency','Status','Created at'])
    for r in rows:
        ws.append([r['id'], r['user_id'], r['fio'], r['program'], r['amount'], r['currency'], r['status'], r['created_at']])
    bio=BytesIO(); wb.save(bio); bio.seek(0)
    resp=make_response(bio.read()); resp.headers['Content-Type']='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'; resp.headers['Content-Disposition']='attachment; filename=payments.xlsx'
    return resp

@app.route('/admin/export/docx')
@role_required('admin')
def export_docx():
    conn=get_db()
    total=conn.execute("SELECT COUNT(*) c FROM payments").fetchone()['c']
    paid=conn.execute("SELECT COUNT(*) c FROM payments WHERE status='paid'").fetchone()['c']
    summ=conn.execute("SELECT SUM(amount) s FROM payments WHERE status='paid'").fetchone()['s'] or 0
    conn.close()
    doc=Document(); doc.add_heading('Отчёт по оплате обучения', level=1)
    doc.add_paragraph('ЧОУВО «Московский университет имени С.Ю. Витте»')
    doc.add_paragraph(f'Всего платежей: {total}'); doc.add_paragraph(f'Оплачено: {paid}'); doc.add_paragraph(f'Сумма: {summ:.2f}')
    fname=f'report_{int(datetime.now().timestamp())}.docx'; fpath=os.path.join(UPLOAD_FOLDER,fname); doc.save(fpath)
    return send_from_directory(UPLOAD_FOLDER, fname, as_attachment=True)

@app.errorhandler(404)
def not_found(e): return render_template('404.html'), 404

@app.errorhandler(500)
def err500(e): return render_template('500.html'), 500

if __name__=='__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)
